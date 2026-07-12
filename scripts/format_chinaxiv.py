from __future__ import annotations

import csv
import re
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
PAPER_MD = ROOT / "paper" / "paper.md"
BIB = ROOT / "paper" / "references.bib"
MODEL_TABLE = ROOT / "paper" / "tables" / "model_taxonomy.csv"
BENCHMARK_TABLE = ROOT / "paper" / "tables" / "benchmark_taxonomy.csv"
MODEL_COMPARISON_TABLE = ROOT / "paper" / "tables" / "model_comparison.csv"
MODEL_FIGURE = ROOT / "paper" / "figures" / "model_timeline.png"
BENCHMARK_FIGURE = ROOT / "paper" / "figures" / "benchmark_timeline.png"
ARCHITECTURE_FIGURE = ROOT / "paper" / "figures" / "architecture_diagram.png"
CONNECTOR_FIGURE = ROOT / "paper" / "figures" / "connector_tradeoff_diagram.png"
BENCHMARK_TAXONOMY_FIGURE = ROOT / "paper" / "figures" / "benchmark_taxonomy_diagram.png"
GROUNDING_FIGURE = ROOT / "paper" / "figures" / "grounding_workflow_diagram.png"
HALLUCINATION_FIGURE = ROOT / "paper" / "figures" / "hallucination_illustration.png"
OUT_MD = ROOT / "paper" / "paper_chinaxiv.md"
OUT_DOCX = ROOT / "paper" / "paper_chinaxiv.docx"


CITATION_RE = re.compile(r"\[@([A-Za-z0-9:_-]+)\]")
BIB_ENTRY_RE = re.compile(r"@(\w+)\{([^,]+),([\s\S]*?)\n\}", re.MULTILINE)
BIB_FIELD_RE = re.compile(r"\n\s*(\w+)\s*=\s*\{([\s\S]*?)\},?", re.MULTILINE)


def parse_bib(path: Path) -> dict[str, dict[str, str]]:
    text = path.read_text(encoding="utf-8")
    entries: dict[str, dict[str, str]] = {}
    for entry_type, key, body in BIB_ENTRY_RE.findall(text):
        fields = {"ENTRYTYPE": entry_type}
        for field, value in BIB_FIELD_RE.findall(body):
            fields[field.lower()] = " ".join(value.split())
        entries[key] = fields
    return entries


def citation_order(text: str) -> OrderedDict[str, int]:
    order: OrderedDict[str, int] = OrderedDict()
    for key in CITATION_RE.findall(text):
        if key not in order:
            order[key] = len(order) + 1
    return order


def replace_citations(text: str, order: OrderedDict[str, int]) -> str:
    def repl(match: re.Match[str]) -> str:
        return f"[{order[match.group(1)]}]"

    return CITATION_RE.sub(repl, text)


def split_authors(author_field: str) -> list[str]:
    if not author_field:
        return []
    return [a.strip() for a in author_field.split(" and ")]


def normalize_author(author: str) -> str:
    author = author.strip("{}")
    if "," in author:
        last, first = [part.strip() for part in author.split(",", 1)]
        return f"{last} {first}"
    return author


def format_authors(author_field: str) -> str:
    authors = [normalize_author(a) for a in split_authors(author_field)]
    if not authors:
        return ""
    if len(authors) > 3:
        return ", ".join(authors[:3]) + ", et al"
    return ", ".join(authors)


def format_reference(index: int, entry: dict[str, str]) -> str:
    authors = format_authors(entry.get("author", ""))
    title = entry.get("title", "").rstrip(".")
    year = entry.get("year", "")
    journal = entry.get("journal", "")
    booktitle = entry.get("booktitle", "")
    institution = entry.get("institution", "")
    url = entry.get("url", "")
    entry_type = entry.get("ENTRYTYPE", "").lower()
    author_part = authors.rstrip(".")

    if entry_type == "inproceedings":
        source = booktitle
        return f"[{index}] {author_part}. {title}[C]. {source}, {year}."
    if entry_type == "article":
        source = journal
        return f"[{index}] {author_part}. {title}[J]. {source}, {year}."
    if entry_type == "techreport":
        suffix = f" {url}" if url else ""
        return f"[{index}] {author_part}. {title}[R]. {institution}, {year}.{suffix}"
    if entry_type == "misc":
        source = entry.get("howpublished", "")
        suffix = f" {url}" if url else ""
        return f"[{index}] {author_part}. {title}[Z]. {source}, {year}.{suffix}"
    return f"[{index}] {author_part}. {title}. {year}."


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def model_summary_rows() -> list[dict[str, str]]:
    return [
        {
            "stage": "Pre-LLM VLP",
            "examples": "ViLBERT, LXMERT, UNITER, Oscar, VinVL",
            "main pattern": "cross-modal encoder pretraining",
            "role in survey": "foundation for image-text representation learning",
        },
        {
            "stage": "Open-vocabulary alignment",
            "examples": "CLIP, ALIGN, OFA, PaLI",
            "main pattern": "contrastive or sequence-to-sequence scaling",
            "role in survey": "transferable visual-language representations",
        },
        {
            "stage": "Frozen-connector MLLMs",
            "examples": "Flamingo, BLIP-2, MiniGPT-4",
            "main pattern": "vision encoder connected to frozen or partly frozen LLM",
            "role in survey": "efficient multimodal generation and few-shot adaptation",
        },
        {
            "stage": "Instruction-tuned MLLMs",
            "examples": "InstructBLIP, LLaVA, Qwen-VL, mPLUG-Owl",
            "main pattern": "visual instruction tuning and alignment",
            "role in survey": "assistant-style multimodal interaction",
        },
        {
            "stage": "Grounded and visual-expert MLLMs",
            "examples": "Kosmos-2, Shikra, Ferret, CogVLM, InternVL",
            "main pattern": "grounding, localization, expert modules, and scale",
            "role in survey": "fine-grained visual evidence and grounding reliability",
        },
        {
            "stage": "Expanded modality systems",
            "examples": "PaLM-E, Video-LLaVA, Visual ChatGPT, GPT-4V",
            "main pattern": "embodiment, video, tool use, or closed frontier alignment",
            "role in survey": "broader deployment and reproducibility trade-offs",
        },
    ]


def benchmark_summary_rows() -> list[dict[str, str]]:
    return [
        {
            "benchmark group": "General VQA",
            "examples": "VQA, GQA, OK-VQA",
            "primary capability": "image-conditioned question answering",
            "evaluation risk": "language priors and dataset bias",
        },
        {
            "benchmark group": "OCR and chart reasoning",
            "examples": "TextVQA, ChartQA",
            "primary capability": "reading scene text and structured graphics",
            "evaluation risk": "small text, chart parsing, and answer normalization",
        },
        {
            "benchmark group": "Science and mathematics",
            "examples": "ScienceQA, MathVista",
            "primary capability": "diagram-grounded and visual mathematical reasoning",
            "evaluation risk": "fluent but unsupported reasoning",
        },
        {
            "benchmark group": "Broad diagnostics",
            "examples": "MME, MMBench, MM-Vet",
            "primary capability": "perception, cognition, and integrated skills",
            "evaluation risk": "prompt sensitivity and scoring differences",
        },
        {
            "benchmark group": "Hallucination checks",
            "examples": "POPE",
            "primary capability": "detecting unsupported object claims",
            "evaluation risk": "coverage beyond object presence remains limited",
        },
        {
            "benchmark group": "Expert-domain reasoning",
            "examples": "MMMU",
            "primary capability": "college-level multidisciplinary multimodal reasoning",
            "evaluation risk": "contamination and domain coverage",
        },
    ]


def display_rows(path: Path) -> list[dict[str, str]]:
    if path == MODEL_TABLE:
        return model_summary_rows()
    if path == BENCHMARK_TABLE:
        return benchmark_summary_rows()
    return read_csv(path)


def csv_to_markdown(path: Path) -> str:
    rows = display_rows(path)
    headers = list(rows[0].keys())
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row[h] for h in headers) + " |")
    return "\n".join(lines)


def csv_to_markdown_columns(path: Path, headers: list[str]) -> str:
    rows = read_csv(path)
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        lines.append("| " + " | ".join(row[h] for h in headers) + " |")
    return "\n".join(lines)


def build_chinaxiv_markdown() -> tuple[str, list[str], list[str]]:
    original = PAPER_MD.read_text(encoding="utf-8")
    bib = parse_bib(BIB)
    order = citation_order(original)
    text = replace_citations(original, order)

    text = text.replace("## References\n\nThe bibliography is maintained in `references.bib`.\n", "").rstrip()

    architecture_figure = "Figure 1. General MLLM architecture and information-flow bottleneck.\n\n![Figure 1](figures/architecture_diagram.png)"
    model_comparison_part_a_headers = ["Model", "Year", "Vision Encoder", "LLM Backbone", "Connector", "Open Source"]
    model_comparison_part_b_headers = [
        "Model",
        "Complexity and Resource Demand",
        "Applicable Scenario",
        "Major Strength",
        "Major Weakness",
    ]
    comparison_table = (
        "Table 1a. Representative MLLM architecture and openness comparison.\n\n"
        + csv_to_markdown_columns(MODEL_COMPARISON_TABLE, model_comparison_part_a_headers)
        + "\n\n"
        + "Table 1b. Representative MLLM resource demand, scenario, and limitation comparison.\n\n"
        + csv_to_markdown_columns(MODEL_COMPARISON_TABLE, model_comparison_part_b_headers)
    )
    model_table = "Table 2. Representative MLLM architecture taxonomy.\n\n" + csv_to_markdown(MODEL_TABLE)
    connector_figure = "Figure 2. Connector trade-offs in MLLM design.\n\n![Figure 2](figures/connector_tradeoff_diagram.png)"
    model_figure = "Figure 3. Timeline distribution of representative MLLM-related models.\n\n![Figure 3](figures/model_timeline.png)"
    benchmark_table = "Table 3. Representative MLLM benchmark taxonomy.\n\n" + csv_to_markdown(BENCHMARK_TABLE)
    benchmark_taxonomy_figure = "Figure 4. Benchmark taxonomy for MLLM evaluation.\n\n![Figure 4](figures/benchmark_taxonomy_diagram.png)"
    benchmark_figure = "Figure 5. Timeline distribution of representative MLLM benchmarks.\n\n![Figure 5](figures/benchmark_timeline.png)"
    hallucination_figure = "Figure 6. Illustrative hallucination mechanism.\n\n![Figure 6](figures/hallucination_illustration.png)"
    grounding_figure = "Figure 7. Evidence-grounded response workflow.\n\n![Figure 7](figures/grounding_workflow_diagram.png)"

    text = text.replace(
        "### 3.2 Survey Methodology",
        architecture_figure + "\n\n### 3.2 Survey Methodology",
        1,
    )

    text = text.replace(
        "### 4.7 Training Paradigms",
        comparison_table + "\n\n" + model_table + "\n\n" + connector_figure + "\n\n" + model_figure + "\n\n### 4.7 Training Paradigms",
        1,
    )
    text = text.replace(
        "### 4.8.6 Comparative Analysis Without Fabricated Scores",
        benchmark_table + "\n\n" + benchmark_taxonomy_figure + "\n\n" + benchmark_figure + "\n\n### 4.8.6 Comparative Analysis Without Fabricated Scores",
        1,
    )
    text = text.replace(
        "The scientifically preferable answer would be:",
        hallucination_figure + "\n\nThe scientifically preferable answer would be:",
        1,
    )
    text = text.replace(
        "The correct system behavior should include uncertainty:",
        grounding_figure + "\n\nThe correct system behavior should include uncertainty:",
        1,
    )

    references = [format_reference(number, bib[key]) for key, number in order.items()]
    text += "\n\n## References\n\n" + "\n\n".join(references) + "\n"
    return text, references, list(order.keys())


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_top_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), "000000")
    p_bdr.append(top)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 6) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text.replace("_", " "))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)


def add_csv_table(
    doc: Document,
    caption: str,
    path: Path,
    headers: list[str] | None = None,
    font_size: int = 6,
) -> None:
    para = doc.add_paragraph(caption)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.bold = True

    rows = display_rows(path) if headers is None else read_csv(path)
    headers = list(rows[0].keys()) if headers is None else headers
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=font_size)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, header in enumerate(headers):
            set_cell_text(cells[idx], row[header], size=font_size)
    source = doc.add_paragraph("Source: Compiled by the author based on cited literature and course materials.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in source.runs:
        set_normal_run(run, size=8)


def add_figure(doc: Document, caption: str, image_path: Path) -> None:
    if image_path.exists():
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(image_path), width=Cm(11.8))
    para = doc.add_paragraph(caption)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in para.runs:
        set_normal_run(run, size=8, bold=True)
    source = doc.add_paragraph("Source: Generated by the author using repository scripts from structured survey metadata.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in source.runs:
        set_normal_run(run, size=8)


def set_normal_run(run, size: int = 10, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc: Document, text: str, style: str | None = None, align=None) -> None:
    para = doc.add_paragraph(style=style)
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_normal_run(run)


def add_heading(doc: Document, text: str, level: int) -> None:
    para = doc.add_paragraph()
    if level == 1:
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        size = 12
    else:
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        size = 10
    run = para.add_run(text)
    set_normal_run(run, size=size, bold=True)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_top_border(paragraph)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_normal_run(run, size=9)


def create_docx(china_md: str, references: list[str]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.footer_distance = Cm(1.2)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    styles["Normal"].font.size = Pt(10)

    lines = china_md.splitlines()
    title = lines[0].lstrip("# ").strip()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_normal_run(r, size=14, bold=True)

    metadata = [
        "ISSA ISSA RASHID",
        "Student ID: 25SF51115",
        "Supervisor / Course Instructor: Dr. Kuo-Kun Tseng",
        "School of Computer Science and Technology, Harbin Institute of Technology, Shenzhen, China",
        "Advanced Artificial Intelligence Course Laboratory",
        "Corresponding author: ISSA ISSA RASHID, 25sf51115@stu.hit.edu.cn",
        "ChinaXiv preprint style manuscript",
    ]
    for item in metadata:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(item)
        set_normal_run(r, size=10)

    doc.add_paragraph()

    in_references = False
    in_code_block = False
    idx = 1
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()
        idx += 1

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue
        if (
            stripped.startswith("Authors:")
            or stripped.startswith("Student ID:")
            or stripped.startswith("Affiliation:")
            or stripped.startswith("Supervisor / Course Instructor:")
            or stripped.startswith("Laboratory:")
            or stripped.startswith("Corresponding author:")
        ):
            continue
        if stripped.startswith("Table 1a."):
            add_csv_table(
                doc,
                stripped,
                MODEL_COMPARISON_TABLE,
                headers=["Model", "Year", "Vision Encoder", "LLM Backbone", "Connector", "Open Source"],
                font_size=6,
            )
            continue
        if stripped.startswith("Table 1b."):
            add_csv_table(
                doc,
                stripped,
                MODEL_COMPARISON_TABLE,
                headers=[
                    "Model",
                    "Complexity and Resource Demand",
                    "Applicable Scenario",
                    "Major Strength",
                    "Major Weakness",
                ],
                font_size=6,
            )
            continue
        if stripped.startswith("Table 2. Representative MLLM architecture"):
            add_csv_table(doc, stripped, MODEL_TABLE, font_size=7)
            continue
        if stripped.startswith("Table 3. Representative MLLM benchmark"):
            add_csv_table(doc, stripped, BENCHMARK_TABLE, font_size=7)
            continue
        figure_map = {
            "Figure 1.": ARCHITECTURE_FIGURE,
            "Figure 2.": CONNECTOR_FIGURE,
            "Figure 3.": MODEL_FIGURE,
            "Figure 4.": BENCHMARK_TAXONOMY_FIGURE,
            "Figure 5.": BENCHMARK_FIGURE,
            "Figure 6.": HALLUCINATION_FIGURE,
            "Figure 7.": GROUNDING_FIGURE,
        }
        matched_figure = next((path for prefix, path in figure_map.items() if stripped.startswith(prefix)), None)
        if matched_figure is not None:
            add_figure(doc, stripped, matched_figure)
            continue
        if stripped.startswith("!["):
            continue
        if stripped.startswith("|"):
            continue
        if not stripped:
            continue
        if stripped == "## References":
            in_references = True
            add_heading(doc, "References", 1)
            continue
        if in_references:
            add_paragraph(doc, stripped)
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:], 1)
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:], 2)
            continue
        if stripped.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(stripped[2:])
            set_normal_run(run)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            run = para.add_run(re.sub(r"^\d+\.\s+", "", stripped))
            set_normal_run(run)
            continue
        add_paragraph(doc, stripped)

    doc.save(OUT_DOCX)


def main() -> None:
    china_md, references, _keys = build_chinaxiv_markdown()
    OUT_MD.write_text(china_md, encoding="utf-8")
    create_docx(china_md, references)
    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()
